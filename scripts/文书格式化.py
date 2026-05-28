#!/usr/bin/env python3
"""
检察文书格式化脚本
使用 python-docx 生成符合公文格式的 Word 检察文书

公文格式要求：
- 字体：方正仿宋_GBK（正文三号），方正小标宋_GBK（标题二号），方正楷体_GBK（引用三号）
- 行距：固定值28磅
- 页边距：上37mm，下35mm，左28mm，右26mm
- 页码：居中，四号半角阿拉伯数字
"""

import argparse
import json
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要安装 python-docx 库")
    print("请运行：pip install python-docx")
    sys.exit(1)


def set_font(run, font_name="方正仿宋_GBK", font_size=Pt(16), bold=False):
    """设置字体"""
    run.font.size = font_size
    run.font.bold = bold
    run.font.name = font_name
    # 设置中文字体
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_format(paragraph, line_spacing=Pt(28), alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                         first_line_indent=None, space_before=Pt(0), space_after=Pt(0)):
    """设置段落格式"""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.alignment = alignment
    pf.space_before = space_before
    pf.space_after = space_after
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent


def set_page_margins(document, top=Mm(37), bottom=Mm(35), left=Mm(28), right=Mm(26)):
    """设置页边距"""
    section = document.sections[0]
    section.top_margin = top
    section.bottom_margin = bottom
    section.left_margin = left
    section.right_margin = right


def add_title(document, text, font_name="方正小标宋_GBK", font_size=Pt(22)):
    """添加标题（居中，二号字）"""
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                         first_line_indent=None)
    run = paragraph.add_run(text)
    set_font(run, font_name=font_name, font_size=font_size, bold=False)
    return paragraph


def add_heading_text(document, text, font_name="方正黑体_GBK", font_size=Pt(16)):
    """添加小标题（三号黑体）"""
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0.74))  # 首行缩进两字符
    run = paragraph.add_run(text)
    set_font(run, font_name=font_name, font_size=font_size, bold=True)
    return paragraph


def add_body_text(document, text, font_name="方正仿宋_GBK", font_size=Pt(16),
                  indent=True, bold=False):
    """添加正文段落（三号仿宋，首行缩进两字符）"""
    paragraph = document.add_paragraph()
    indent_val = Cm(0.74) if indent else None  # 三号字约0.74cm缩进两字符
    set_paragraph_format(paragraph, first_line_indent=indent_val)
    run = paragraph.add_run(text)
    set_font(run, font_name=font_name, font_size=font_size, bold=bold)
    return paragraph


def add_reference_text(document, text, font_name="方正楷体_GBK", font_size=Pt(16)):
    """添加引用文字（三号楷体）"""
    paragraph = document.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=Cm(0.74))
    run = paragraph.add_run(text)
    set_font(run, font_name=font_name, font_size=font_size)
    return paragraph


def add_empty_line(document, count=1):
    """添加空行"""
    for _ in range(count):
        document.add_paragraph()


def create_qishu(data):
    """
    生成起诉书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称（如"北京市海淀区人民检察院"）
    - case_number: 文号（如"京海检刑诉〔2024〕123号"）
    - defendant: 被告人信息 dict (name, gender, birth_date, id_number, ethnicity,
                                    education, household, address, occupation)
    - arrest_date: 被拘留日期
    - arrest_approval_date: 批准逮捕日期
    - arrest_execution_date: 执行逮捕日期
    - detention_place: 羁押场所
    - facts: 案件事实 (str)
    - evidence: 证据列表 (list of str)
    - legal_basis: 法律依据 (str)
    - sentencing_circumstances: 量刑情节 (str)
    - plea_agreement: 认罪认罚情况 (str or None)
    - sentencing_recommendation: 量刑建议 (str)
    """
    doc = Document()
    set_page_margins(doc)

    # 机关名称
    add_title(doc, data.get("procuratorate", "×××人民检察院"),
              font_name="方正小标宋_GBK", font_size=Pt(22))

    # 文书名称
    add_title(doc, "起 诉 书",
              font_name="方正小标宋_GBK", font_size=Pt(22))

    # 文号
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×刑诉〔20××〕×号"))
    set_font(run, font_size=Pt(16))

    add_empty_line(doc)

    # 被告人基本情况
    add_heading_text(doc, "一、被告人基本情况")
    d = data.get("defendant", {})
    gender = d.get("gender", "男/女")
    defendant_info = (
        f"被告人{d.get('name', '×××')}，{gender}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"户籍所在地{d.get('household', '×××')}，"
        f"现住{d.get('address', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    if data.get("arrest_date"):
        defendant_info += f"因涉嫌×××罪，于{data['arrest_date']}被×××公安局刑事拘留，"
    if data.get("arrest_approval_date"):
        defendant_info += f"于{data['arrest_approval_date']}经本院批准逮捕，"
    if data.get("arrest_execution_date"):
        defendant_info += f"于{data['arrest_execution_date']}由×××公安局执行逮捕，"
    defendant_info += f"现羁押于{data.get('detention_place', '×××看守所')}。"
    add_body_text(doc, defendant_info)

    add_empty_line(doc)

    # 案件事实
    add_heading_text(doc, "二、案件事实")
    add_body_text(doc, "经依法审查查明：")
    add_body_text(doc, data.get("facts", "（案件事实）"))

    add_empty_line(doc)

    # 证据
    add_heading_text(doc, "认定上述事实的证据如下：")
    for i, evidence in enumerate(data.get("evidence", []), 1):
        add_body_text(doc, f"{i}. {evidence}")

    add_empty_line(doc)

    # 起诉理由
    add_heading_text(doc, "三、起诉理由和法律依据")
    legal_basis = data.get("legal_basis", "×××罪")
    add_body_text(
        doc,
        f"本院认为，被告人×××的行为触犯了{legal_basis}，"
        f"犯罪事实清楚，证据确实、充分，应当以×××罪追究其刑事责任。"
    )

    # 量刑情节
    if data.get("sentencing_circumstances"):
        add_body_text(doc, "量刑情节分析：")
        add_body_text(doc, data["sentencing_circumstances"])

    # 认罪认罚
    if data.get("plea_agreement"):
        add_body_text(doc, f"认罪认罚情况：{data['plea_agreement']}")

    # 量刑建议
    add_body_text(doc, f"量刑建议：{data.get('sentencing_recommendation', '（量刑建议）')}")

    add_body_text(doc, "根据《中华人民共和国刑事诉讼法》第一百七十六条的规定提起公诉。")

    add_empty_line(doc, 2)

    # 此致
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0.74))
    run = p.add_run("此致")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                         first_line_indent=Cm(0.74))
    run = p.add_run("×××人民法院")
    set_font(run)

    add_empty_line(doc, 2)

    # 签名
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("检察员：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_buqisu(data):
    """
    生成不起诉决定书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - unprosecuted_person: 被不起诉人信息 (同 defendant)
    - facts: 案件事实
    - unprosecution_type: 不起诉类型 ("法定" / "酌定" / "存疑" / "附条件")
    - reasons: 不起诉理由
    - legal_basis: 法律依据
    - victim_info: 被害人信息（用于告知申诉权利）
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "不 起 诉 决 定 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×刑不诉〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    # 被不起诉人基本情况
    add_heading_text(doc, "一、被不起诉人基本情况")
    d = data.get("unprosecuted_person", {})
    gender = d.get("gender", "男/女")
    info = (
        f"被不起诉人{d.get('name', '×××')}，{gender}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"户籍所在地{d.get('household', '×××')}，"
        f"现住{d.get('address', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    # 案件事实
    add_heading_text(doc, "二、案件事实")
    add_body_text(doc, data.get("facts", "（案件事实）"))

    add_empty_line(doc)

    # 不起诉理由
    add_heading_text(doc, "三、不起诉理由和法律依据")
    unprosecution_type = data.get("unprosecution_type", "酌定")

    if unprosecution_type == "法定":
        add_body_text(doc,
            "本院认为，被不起诉人没有犯罪事实/具有《中华人民共和国刑事诉讼法》"
            "第十六条规定的情形之一。")
        add_body_text(doc,
            "根据《中华人民共和国刑事诉讼法》第一百七十七条第一款的规定，"
            "决定对被不起诉人不起诉。")
    elif unprosecution_type == "酌定":
        add_body_text(doc,
            f"本院认为，被不起诉人的行为触犯了{data.get('legal_basis', '×××')}"
            "，构成×××罪。但鉴于犯罪情节轻微，"
            f"{data.get('reasons', '不需要判处刑罚')}。")
        add_body_text(doc,
            "根据《中华人民共和国刑事诉讼法》第一百七十七条第二款的规定，"
            "决定对被不起诉人不起诉。")
    elif unprosecution_type == "存疑":
        add_body_text(doc, "本院认为，本案证据不足，不符合起诉条件。")
        add_body_text(doc,
            "根据《中华人民共和国刑事诉讼法》第一百七十五条第四款的规定，"
            "决定对被不起诉人不起诉。")
    elif unprosecution_type == "附条件":
        add_body_text(doc,
            "被不起诉人（未成年人）实施了×××行为，涉嫌×××罪，"
            "可能判处一年有期徒刑以下刑罚，符合起诉条件，但有悔罪表现。")
        add_body_text(doc,
            "根据《中华人民共和国刑事诉讼法》第二百八十二条第一款的规定，"
            "决定对被不起诉人附条件不起诉。")

    add_empty_line(doc, 2)

    # 签名
    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("检察员：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_jianchayijian(data):
    """
    生成审查逮捕意见书

    data: dict
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "审 查 逮 捕 意 见 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×捕审〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    # 案件来源
    add_heading_text(doc, "一、案件来源")
    add_body_text(doc, data.get("case_source",
        "犯罪嫌疑人×××涉嫌×××罪一案，由×××公安局提请批准逮捕。"))

    add_empty_line(doc)

    # 犯罪嫌疑人基本情况
    add_heading_text(doc, "二、犯罪嫌疑人基本情况")
    d = data.get("suspect", {})
    info = (
        f"犯罪嫌疑人{d.get('name', '×××')}，{d.get('gender', '男')}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"户籍所在地{d.get('household', '×××')}，"
        f"现住{d.get('address', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    # 案件事实
    add_heading_text(doc, "三、案件事实")
    add_body_text(doc, data.get("facts", "（案件事实）"))

    add_empty_line(doc)

    # 证据情况
    add_heading_text(doc, "四、证据情况")
    for i, evidence in enumerate(data.get("evidence", []), 1):
        add_body_text(doc, f"{i}. {evidence}")

    add_empty_line(doc)

    # 审查意见
    add_heading_text(doc, "五、审查意见")

    add_body_text(doc, "（一）犯罪构成分析", bold=True)
    add_body_text(doc, data.get("composition_analysis", "（犯罪构成分析）"))

    add_body_text(doc, "（二）证据审查", bold=True)
    add_body_text(doc, data.get("evidence_review", "（证据审查意见）"))

    add_body_text(doc, "（三）逮捕必要性审查", bold=True)
    add_body_text(doc, data.get("necessity_review", "（社会危险性评估）"))

    add_body_text(doc, "（四）认罪认罚情况", bold=True)
    add_body_text(doc, data.get("plea_status", "（认罪认罚情况）"))

    add_empty_line(doc)

    # 处理意见
    add_heading_text(doc, "六、处理意见")
    opinion = data.get("opinion",
        "综上所述，犯罪嫌疑人×××的行为涉嫌×××罪，事实清楚，证据确实、充分，"
        "符合逮捕条件。本院拟作出批准逮捕的决定。")
    add_body_text(doc, opinion)

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("承办人：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_kangsu(data):
    """
    生成抗诉书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - original_judgment: 原审判决/裁定概况
    - defendant: 被告人信息
    - reasons: 抗诉理由列表 [{"type": "事实认定错误", "content": "..."}]
    - legal_basis: 法律依据
    - opinion: 抗诉意见
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "抗 诉 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×刑抗〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    add_heading_text(doc, "一、原审判决/裁定概况")
    add_body_text(doc, data.get("original_judgment", "（原审判决/裁定概况）"))

    add_empty_line(doc)

    add_heading_text(doc, "二、被告人基本情况")
    d = data.get("defendant", {})
    info = (
        f"被告人{d.get('name', '×××')}，{d.get('gender', '男')}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"户籍所在地{d.get('household', '×××')}，"
        f"现住{d.get('address', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    add_heading_text(doc, "三、抗诉理由")
    add_body_text(doc, "本院认为，上述判决/裁定确有错误，理由如下：")

    for i, reason in enumerate(data.get("reasons", []), 1):
        add_body_text(doc, f"（{_cn_num(i)}）{reason.get('type', '×××错误')}", bold=True)
        add_body_text(doc, reason.get("content", "（具体理由）"))

    add_empty_line(doc)

    add_heading_text(doc, "四、法律依据")
    add_body_text(doc, data.get("legal_basis",
        "根据《中华人民共和国刑事诉讼法》第二百二十八条之规定，"
        "特向×××人民法院提出抗诉。"))

    add_empty_line(doc)

    add_heading_text(doc, "五、抗诉意见")
    add_body_text(doc, data.get("opinion",
        "综上所述，原审判决/裁定存在上述错误，请依法改判/发回重审。"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("court", "×××人民法院"))
    set_font(run)

    add_empty_line(doc)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("procuratorate", "×××人民检察院"))
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("检察员：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_shencha_baogao(data):
    """
    生成案件审查报告

    data: dict
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, "案 件 审 查 报 告")

    add_empty_line(doc)

    # 案件基本信息
    add_heading_text(doc, "一、案件来源")
    add_body_text(doc, data.get("case_source",
        "×××公安局于20××年×月×日将犯罪嫌疑人×××涉嫌×××罪一案移送本院审查起诉。"))

    add_empty_line(doc)

    add_heading_text(doc, "二、犯罪嫌疑人/被告人基本情况")
    for i, d in enumerate(data.get("defendants", [data.get("defendant", {})]), 1):
        info = (
            f"{i}. 被告人{d.get('name', '×××')}，{d.get('gender', '男')}，"
            f"{d.get('birth_date', '×年×月×日')}出生，"
            f"公民身份号码{d.get('id_number', '×××')}，"
            f"{d.get('ethnicity', '×')}族，"
            f"文化程度{d.get('education', '×××')}，"
            f"职业{d.get('occupation', '×××')}。"
        )
        add_body_text(doc, info)

    add_empty_line(doc)

    add_heading_text(doc, "三、案件事实")
    add_body_text(doc, "（一）公安机关认定的犯罪事实", bold=True)
    add_body_text(doc, data.get("police_facts", "（摘录起诉意见书认定的犯罪事实）"))
    add_body_text(doc, "（二）经审查查明的犯罪事实", bold=True)
    add_body_text(doc, data.get("facts", "（根据证据审查后认定的事实）"))

    add_empty_line(doc)

    add_heading_text(doc, "四、证据情况")
    evidence = data.get("evidence", [])
    if evidence:
        for i, ev in enumerate(evidence, 1):
            add_body_text(doc, f"{i}. {ev}")
    else:
        add_body_text(doc, "（证据清单）")

    add_body_text(doc, "证据分析：", bold=True)
    add_body_text(doc, data.get("evidence_analysis", "（证据合法性、真实性、关联性分析）"))

    add_empty_line(doc)

    add_heading_text(doc, "五、法律适用分析")
    add_body_text(doc, data.get("legal_analysis", "（犯罪构成分析、此罪彼罪分析）"))

    add_empty_line(doc)

    add_heading_text(doc, "六、量刑情节分析")
    add_body_text(doc, data.get("sentencing_analysis", "（法定和酌定量刑情节分析）"))

    add_empty_line(doc)

    add_heading_text(doc, "七、认罪认罚情况")
    add_body_text(doc, data.get("plea_status", "（认罪认罚情况说明）"))

    add_empty_line(doc)

    add_heading_text(doc, "八、处理意见")
    add_body_text(doc, data.get("opinion",
        "经审查，本案犯罪事实清楚，证据确实、充分，建议提起公诉。"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("承办人：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_jiancha_jianyi(data):
    """
    生成检察建议书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - target_unit: 被建议单位名称
    - case_source: 案件来源
    - facts: 调查核实情况
    - problems: 存在的问题列表
    - causes: 问题原因分析
    - suggestions: 检察建议列表
    - legal_basis: 法律依据
    - deadline: 落实期限（默认"三个月"）
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "检 察 建 议 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×建〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    # 被建议单位
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0.74))
    run = p.add_run(f"{data.get('target_unit', '×××（被建议单位）')}：")
    set_font(run)

    # 案件来源
    add_heading_text(doc, "一、案件来源")
    add_body_text(doc, data.get("case_source",
        "本院在办理×××案件/开展×××工作中，发现以下问题："))

    add_empty_line(doc)

    # 调查核实情况
    add_heading_text(doc, "二、调查核实情况")
    add_body_text(doc, data.get("facts", "（调查核实过程和发现的事实）"))

    add_empty_line(doc)

    # 存在的问题
    add_heading_text(doc, "三、存在的问题")
    for i, problem in enumerate(data.get("problems", ["（问题一）"]), 1):
        add_body_text(doc, f"{i}. {problem}")

    add_empty_line(doc)

    # 问题原因
    if data.get("causes"):
        add_heading_text(doc, "四、问题分析")
        add_body_text(doc, data["causes"])

        add_empty_line(doc)
        sug_section = "五、检察建议"
    else:
        sug_section = "四、检察建议"

    # 检察建议
    add_heading_text(doc, sug_section)
    add_body_text(doc,
        "为促进依法行政/严格执法/公正司法/社会治理，"
        "根据《中华人民共和国人民检察院组织法》第二十一条之规定，现提出如下建议：")
    for i, sug in enumerate(data.get("suggestions", ["（建议一）"]), 1):
        add_body_text(doc, f"{i}. {sug}")

    add_empty_line(doc)

    # 法律依据
    add_heading_text(doc, "法律依据")
    add_body_text(doc, data.get("legal_basis", "（法律依据）"))

    add_body_text(doc,
        f"请你单位自收到本建议书之日起{data.get('deadline', '三个月')}内依法落实，"
        "并将落实情况书面回复本院。")

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("procuratorate", "×××人民检察院"))
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_jiuzheng(data):
    """
    生成纠正违法通知书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - target_unit: 被通知单位名称（公安/法院/看守所等）
    - case_source: 案件来源
    - illegal_facts: 违法事实
    - legal_basis: 法律依据
    - correction_opinion: 纠正意见
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "纠 正 违 法 通 知 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×纠违〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    p = doc.add_paragraph()
    set_paragraph_format(p, first_line_indent=Cm(0.74))
    run = p.add_run(f"{data.get('target_unit', '×××')}：")
    set_font(run)

    add_heading_text(doc, "一、案件来源")
    add_body_text(doc, data.get("case_source",
        "本院在办理×××案件/开展法律监督工作中，发现你单位存在以下违法行为："))

    add_empty_line(doc)

    add_heading_text(doc, "二、违法事实")
    add_body_text(doc, data.get("illegal_facts", "（违法事实）"))

    add_empty_line(doc)

    add_heading_text(doc, "三、法律依据")
    add_body_text(doc, data.get("legal_basis", "（法律依据）"))

    add_empty_line(doc)

    add_heading_text(doc, "四、纠正意见")
    add_body_text(doc, data.get("correction_opinion",
        "上述行为违反了法律规定，请你单位依法予以纠正，并将纠正情况书面回复本院。"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("procuratorate", "×××人民检察院"))
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_tuibu_tigang(data):
    """
    生成退回补充侦查提纲

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - case基本情况: 案件基本情况
    - current_evidence: 现有证据情况
    - supplement_items: 需要补充侦查的事项列表
    - supplement_direction: 补充侦查方向和建议
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, "退 回 补 充 侦 查 提 纲")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", ""))
    set_font(run)

    add_empty_line(doc)

    add_heading_text(doc, "一、案件基本情况")
    add_body_text(doc, data.get("case基本情况", "（案件基本情况）"))

    add_empty_line(doc)

    add_heading_text(doc, "二、现有证据情况")
    add_body_text(doc, data.get("current_evidence", "（现有证据情况）"))

    add_empty_line(doc)

    add_heading_text(doc, "三、需要补充侦查的事项")
    for i, item in enumerate(data.get("supplement_items", ["（补充事项一）"]), 1):
        add_body_text(doc, f"{i}. {item}")

    add_empty_line(doc)

    add_heading_text(doc, "四、补充侦查的方向和建议")
    add_body_text(doc, data.get("supplement_direction", "（补充侦查方向和建议）"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("承办人：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_jiya_bishencha(data):
    """
    生成羁押必要性审查建议书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - suspect: 犯罪嫌疑人信息
    - case_facts: 案件事实
    - custody_reason: 羁押理由
    - review_opinion: 审查意见（建议释放/变更强制措施的理由）
    - legal_basis: 法律依据
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "羁 押 必 要 性 审 查 建 议 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", "×检×羁审〔20××〕×号"))
    set_font(run)

    add_empty_line(doc)

    add_heading_text(doc, "一、犯罪嫌疑人基本情况")
    d = data.get("suspect", {})
    info = (
        f"犯罪嫌疑人{d.get('name', '×××')}，{d.get('gender', '男')}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    add_heading_text(doc, "二、案件事实")
    add_body_text(doc, data.get("case_facts", "（案件事实）"))

    add_empty_line(doc)

    add_heading_text(doc, "三、羁押情况")
    add_body_text(doc, data.get("custody_reason", "（羁押理由和现状）"))

    add_empty_line(doc)

    add_heading_text(doc, "四、审查意见")
    add_body_text(doc, data.get("review_opinion", "（审查意见：建议释放或变更强制措施的理由）"))

    add_empty_line(doc)

    add_heading_text(doc, "五、法律依据")
    add_body_text(doc, data.get("legal_basis",
        "根据《中华人民共和国刑事诉讼法》第九十五条之规定，"
        "建议对犯罪嫌疑人×××变更强制措施为取保候审/予以释放。"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("承办人：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def create_renzui_jujieshu(data):
    """
    生成认罪认罚具结书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - suspect: 犯罪嫌疑人信息
    - case_facts: 案件基本情况
    - admitted_facts: 承认的犯罪事实
    - agreed_charges: 同意的罪名
    - sentencing_recommendation: 量刑建议
    - program_suggestion: 程序适用建议（速裁/简易/普通）
    - defense_lawyer: 辩护人/值班律师信息
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "认 罪 认 罚 具 结 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", ""))
    set_font(run)

    add_empty_line(doc)

    # 犯罪嫌疑人基本信息
    add_heading_text(doc, "一、犯罪嫌疑人基本信息")
    d = data.get("suspect", {})
    info = (
        f"犯罪嫌疑人{d.get('name', '×××')}，{d.get('gender', '男')}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    # 案件基本情况
    add_heading_text(doc, "二、案件基本情况")
    add_body_text(doc, data.get("case_facts", "（案件基本情况）"))

    add_empty_line(doc)

    # 认罪认罚内容
    add_heading_text(doc, "三、认罪认罚内容")
    add_body_text(doc, f"犯罪嫌疑人{d.get('name', '×××')}自愿承认以下犯罪事实：")
    add_body_text(doc, data.get("admitted_facts", "（承认的犯罪事实）"))
    add_body_text(doc, f"同意以{data.get('agreed_charges', '×××')}罪追究其刑事责任。")

    add_empty_line(doc)

    # 量刑建议
    add_heading_text(doc, "四、量刑建议")
    sr = data.get("sentencing_recommendation", {})
    if isinstance(sr, dict):
        add_body_text(doc, f"主刑建议：{sr.get('main_penalty', '（主刑）')}")
        add_body_text(doc, f"附加刑建议：{sr.get('additional_penalty', '（附加刑）')}")
        add_body_text(doc, f"缓刑建议：{sr.get('probation', '（是否适用缓刑）')}")
    else:
        add_body_text(doc, str(sr) if sr else "（量刑建议）")

    add_empty_line(doc)

    # 程序适用建议
    add_heading_text(doc, "五、程序适用建议")
    add_body_text(doc, f"建议适用{data.get('program_suggestion', '简易/速裁/普通')}程序。")

    add_empty_line(doc)

    # 法律后果告知
    add_heading_text(doc, "六、法律后果告知")
    add_body_text(doc,
        "犯罪嫌疑人已知悉认罪认罚的法律后果，包括："
        "自愿认罪将依法从宽处理；如违反具结书内容，可能影响从宽幅度。")

    add_empty_line(doc)

    # 辩护人/值班律师
    if data.get("defense_lawyer"):
        add_heading_text(doc, "七、辩护人/值班律师")
        add_body_text(doc, data["defense_lawyer"])
        add_body_text(doc, "辩护人/值班律师确认：犯罪嫌疑人系自愿认罪认罚，已提供法律帮助。")

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(f"犯罪嫌疑人：{d.get('name', '×××')}")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    if data.get("defense_lawyer"):
        add_empty_line(doc)
        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run("辩护人/值班律师：×××")
        set_font(run)

        p = doc.add_paragraph()
        set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(data.get("date", "20××年×月×日"))
        set_font(run)

    return doc


def create_liangxing_jianyi(data):
    """
    生成量刑建议书

    data: dict, 包含以下字段：
    - procuratorate: 机关名称
    - case_number: 文号
    - defendant: 被告人信息
    - charge: 罪名
    - facts_summary: 犯罪事实摘要
    - sentencing_circumstances: 量刑情节分析
    - main_penalty: 主刑建议
    - additional_penalty: 附加刑建议（罚金等）
    - probation: 是否建议缓刑及理由
    - legal_basis: 法律依据
    """
    doc = Document()
    set_page_margins(doc)

    add_title(doc, data.get("procuratorate", "×××人民检察院"))
    add_title(doc, "量 刑 建 议 书")

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(data.get("case_number", ""))
    set_font(run)

    add_empty_line(doc)

    # 被告人基本情况
    add_heading_text(doc, "一、被告人基本情况")
    d = data.get("defendant", {})
    info = (
        f"被告人{d.get('name', '×××')}，{d.get('gender', '男')}，"
        f"{d.get('birth_date', '×年×月×日')}出生，"
        f"公民身份号码{d.get('id_number', '×××')}，"
        f"{d.get('ethnicity', '×')}族，"
        f"文化程度{d.get('education', '×××')}，"
        f"职业{d.get('occupation', '×××')}。"
    )
    add_body_text(doc, info)

    add_empty_line(doc)

    # 涉嫌罪名及犯罪事实
    add_heading_text(doc, "二、涉嫌罪名及犯罪事实")
    add_body_text(doc, f"被告人{d.get('name', '×××')}涉嫌{data.get('charge', '×××')}罪。")
    add_body_text(doc, data.get("facts_summary", "（犯罪事实摘要）"))

    add_empty_line(doc)

    # 量刑情节分析
    add_heading_text(doc, "三、量刑情节分析")
    add_body_text(doc, data.get("sentencing_circumstances", "（量刑情节分析）"))

    add_empty_line(doc)

    # 量刑建议
    add_heading_text(doc, "四、量刑建议")
    add_body_text(doc, f"主刑建议：{data.get('main_penalty', '（主刑建议）')}")
    if data.get("additional_penalty"):
        add_body_text(doc, f"附加刑建议：{data['additional_penalty']}")
    if data.get("probation"):
        add_body_text(doc, f"缓刑建议：{data['probation']}")

    add_empty_line(doc)

    # 法律依据
    add_heading_text(doc, "五、法律依据")
    add_body_text(doc, data.get("legal_basis",
        "根据《中华人民共和国刑事诉讼法》第一百七十六条之规定，"
        "提出以上量刑建议，请合议庭依法采纳。"))

    add_empty_line(doc, 2)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run("检察员：×××")
    set_font(run)

    p = doc.add_paragraph()
    set_paragraph_format(p, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run = p.add_run(data.get("date", "20××年×月×日"))
    set_font(run)

    return doc


def _cn_num(n):
    """数字转中文"""
    cn = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五",
          6: "六", 7: "七", 8: "八", 9: "九", 10: "十"}
    return cn.get(n, str(n))


def main():
    parser = argparse.ArgumentParser(description="检察文书格式化工具")
    type_map = {
        "起诉书": create_qishu,
        "不起诉决定书": create_buqisu,
        "审查逮捕意见书": create_jianchayijian,
        "抗诉书": create_kangsu,
        "案件审查报告": create_shencha_baogao,
        "检察建议书": create_jiancha_jianyi,
        "纠正违法通知书": create_jiuzheng,
        "退补提纲": create_tuibu_tigang,
        "羁押必要性审查建议书": create_jiya_bishencha,
        "认罪认罚具结书": create_renzui_jujieshu,
        "量刑建议书": create_liangxing_jianyi,
    }
    parser.add_argument("type", choices=list(type_map.keys()), help="文书类型")
    parser.add_argument("json_file", help="案件信息 JSON 文件路径")
    parser.add_argument("-o", "--output", help="输出 Word 文件路径（默认与 JSON 文件同名）")

    args = parser.parse_args()

    with open(args.json_file, "r", encoding="utf-8") as f:
        data = json.load(f)
    doc = type_map[args.type](data)

    if args.output:
        output_path = args.output
    else:
        json_path = Path(args.json_file)
        output_path = str(json_path.with_suffix(".docx"))

    doc.save(output_path)
    print(f"文书已生成：{output_path}")


# 供 Claude Code 直接调用的函数接口
def generate_from_dict(doc_type, data, output_path):
    """
    从字典数据生成检察文书

    Args:
        doc_type: 文书类型
        data: 案件信息字典
        output_path: 输出 Word 文件路径

    Returns:
        输出文件路径
    """
    type_map = {
        "起诉书": create_qishu,
        "不起诉决定书": create_buqisu,
        "审查逮捕意见书": create_jianchayijian,
        "抗诉书": create_kangsu,
        "案件审查报告": create_shencha_baogao,
        "检察建议书": create_jiancha_jianyi,
        "纠正违法通知书": create_jiuzheng,
        "退补提纲": create_tuibu_tigang,
        "羁押必要性审查建议书": create_jiya_bishencha,
        "认罪认罚具结书": create_renzui_jujieshu,
        "量刑建议书": create_liangxing_jianyi,
    }
    if doc_type not in type_map:
        raise ValueError(f"不支持的文书类型：{doc_type}，支持：{list(type_map.keys())}")

    doc = type_map[doc_type](data)
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    main()
